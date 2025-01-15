import io
from typing import Any, Dict, List

import docx
import pandas as pd
import PyPDF2
from loguru import logger
from openai import AsyncOpenAI
from sqlalchemy import create_engine
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.order import KnowledgeBase


class KnowledgeBaseService:
    def __init__(self) -> None:
        self.client = AsyncOpenAI(api_key=settings.OPENAI_API_KEY)
        self.embedding_model = "text-embedding-3-small"

    async def process_file(self, file: bytes, filename: str, db: Session) -> None:
        """Process uploaded file and add to knowledge base."""
        try:
            # Extract text based on file type
            text_chunks = await self._extract_text(file, filename)

            # Generate embeddings for each chunk
            embeddings = await self._generate_embeddings(text_chunks)

            # Store in database
            for chunk, embedding in zip(text_chunks, embeddings):
                kb_entry = KnowledgeBase(
                    content=chunk,
                    embedding=embedding,
                    source=filename,
                    metadata={"filename": filename},
                )
                db.add(kb_entry)

            db.commit()
            logger.info(f"Successfully processed file: {filename}")

        except Exception as e:
            logger.error(f"Error processing file {filename}: {str(e)}")
            raise

    async def _extract_text(self, file: bytes, filename: str) -> List[str]:
        """Extract text from various file formats."""
        chunks = []
        file_obj = io.BytesIO(file)

        if filename.endswith((".csv", ".xlsx", ".xls")):
            # Read spreadsheet
            df = (
                pd.read_excel(file_obj)
                if filename.endswith((".xlsx", ".xls"))
                else pd.read_csv(file_obj)
            )
            # Convert each row to a text chunk
            for _, row in df.iterrows():
                chunks.append(" | ".join(f"{col}: {val}" for col, val in row.items()))

        elif filename.endswith(".pdf"):
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(file_obj)
            for page in pdf_reader.pages:
                chunks.append(page.extract_text())

        elif filename.endswith((".doc", ".docx")):
            # Read Word document
            doc = docx.Document(file_obj)
            for para in doc.paragraphs:
                if para.text.strip():
                    chunks.append(para.text)

        else:
            # Try to read as text
            content = file_obj.read().decode("utf-8")
            chunks = content.split("\n\n")

        # Clean and split into smaller chunks if needed
        cleaned_chunks = []
        for chunk in chunks:
            if len(chunk) > 1000:  # Split long chunks
                words = chunk.split()
                for i in range(0, len(words), 200):
                    cleaned_chunks.append(" ".join(words[i : i + 200]))
            else:
                cleaned_chunks.append(chunk)

        return [c.strip() for c in cleaned_chunks if c.strip()]

    async def _generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings for text chunks using OpenAI."""
        try:
            embeddings = []
            for text in texts:
                response = await self.client.embeddings.create(
                    model=self.embedding_model, input=text
                )
                embeddings.append(response.data[0].embedding)
            return embeddings
        except Exception as e:
            logger.error(f"Error generating embeddings: {str(e)}")
            raise

    async def search_similar(
        self, query: str, db: Session, limit: int = 5
    ) -> List[Dict[str, Any]]:
        """Search for similar content using vector similarity."""
        try:
            # Generate query embedding
            query_response = await self.client.embeddings.create(
                model=self.embedding_model, input=query
            )
            query_embedding = query_response.data[0].embedding

            # Perform vector similarity search
            results = (
                db.query(KnowledgeBase)
                .order_by(KnowledgeBase.embedding.cosine_distance(query_embedding))
                .limit(limit)
                .all()
            )

            return [
                {
                    "content": r.content,
                    "source": r.source,
                    "metadata": r.metadata,
                    "created_at": r.created_at,
                }
                for r in results
            ]

        except Exception as e:
            logger.error(f"Error searching knowledge base: {str(e)}")
            return []

    async def connect_database(self, connection_params: Dict[str, str]) -> bool:
        """Connect to external database and import data."""
        try:
            # Create SQLAlchemy engine
            url = f"postgresql://{connection_params['username']}:{connection_params['password']}@{connection_params['host']}:{connection_params['port']}/{connection_params['database']}"
            engine = create_engine(url)

            # Test connection
            with engine.connect() as conn:
                tables = pd.read_sql(
                    "SELECT table_name FROM information_schema.tables WHERE table_schema = 'public'",
                    conn,
                )

                # Process each table
                for table in tables["table_name"]:
                    df = pd.read_sql(f"SELECT * FROM {table}", conn)
                    text_chunks = []

                    # Convert each row to text
                    for _, row in df.iterrows():
                        chunk = f"Table: {table}\n"
                        chunk += "\n".join(f"{col}: {val}" for col, val in row.items())
                        text_chunks.append(chunk)

                    # Generate embeddings and store
                    embeddings = await self._generate_embeddings(text_chunks)

                    # Store in knowledge base
                    for chunk, embedding in zip(text_chunks, embeddings):
                        kb_entry = KnowledgeBase(
                            content=chunk,
                            embedding=embedding,
                            source=f"db:{table}",
                            metadata={
                                "database": connection_params["database"],
                                "table": table,
                            },
                        )
                        Session.object_session(kb_entry).add(kb_entry)

                    Session.object_session(kb_entry).commit()

            return True

        except Exception as e:
            logger.error(f"Error connecting to database: {str(e)}")
            return False
